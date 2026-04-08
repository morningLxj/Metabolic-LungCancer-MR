import os
import math
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = "E:/GWAS"
OUT_DIR = f"{BASE}/Final_Tables_SCI_Format"
MAIN_DOC = f"{OUT_DIR}/Main_Tables_Final.docx"
SUPP_DOC = f"{OUT_DIR}/Supplementary_Tables_Final.docx"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def clear_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "nil"},
                top={"val": "nil"},
                right={"val": "nil"},
                bottom={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )


def apply_three_line(table):
    if len(table.rows) == 0:
        return
    clear_table_borders(table)
    for cell in table.rows[0].cells:
        set_cell_border(
            cell,
            top={"val": "single", "sz": 12, "color": "000000"},
            bottom={"val": "single", "sz": 8, "color": "000000"},
        )
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": 12, "color": "000000"})


def numeric_like(text):
    t = str(text).strip()
    if t == "" or t.lower() in {"na", "n/a", "nan"}:
        return False
    t = t.replace(",", "").replace("%", "").replace("−", "-")
    t0 = t.split(" ")[0]
    try:
        float(t0)
        return True
    except Exception:
        return bool(any(ch.isdigit() for ch in t) and len(t) < 24)


def style_paragraph(p, size=10.5, bold=False, align=None, before=0, after=3):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("Table ") or txt.startswith("**Table") or txt.startswith("Supplementary Table"):
            style_paragraph(p, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=4)
        elif txt.startswith("Module ") or txt.startswith("Supplementary Tables"):
            style_paragraph(p, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)
        else:
            style_paragraph(p, size=10.5, bold=False, align=None, before=0, after=3)
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        for j, cell in enumerate(table.rows[0].cells):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style_paragraph(p, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=1)
        for i in range(1, len(table.rows)):
            row = table.rows[i]
            for j, cell in enumerate(row.cells):
                cell_txt = cell.text
                align = WD_ALIGN_PARAGRAPH.RIGHT if numeric_like(cell_txt) else WD_ALIGN_PARAGRAPH.LEFT
                for p in cell.paragraphs:
                    style_paragraph(p, size=10, bold=False, align=align, before=0, after=0)
        apply_three_line(table)


def add_table_from_df(doc, title, df):
    p = doc.add_paragraph(title)
    style_paragraph(p, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=4)
    if df is None or len(df.columns) == 0:
        q = doc.add_paragraph("No data available")
        style_paragraph(q, size=10, before=0, after=6)
        return
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.autofit = True
    hdr = t.rows[0].cells
    for j, c in enumerate(df.columns):
        hdr[j].text = str(c)
    for _, r in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            v = r[c]
            if isinstance(v, float) and math.isnan(v):
                txt = "NA"
            else:
                txt = str(v).strip()
                if txt == "":
                    txt = "NA"
            cells[j].text = txt
    apply_three_line(t)
    doc.add_paragraph("")


def read_csv(path):
    if os.path.exists(path):
        try:
            return pd.read_csv(path, dtype=str, keep_default_na=False)
        except Exception:
            try:
                return pd.read_csv(path, sep=None, engine="python", dtype=str, keep_default_na=False)
            except Exception:
                try:
                    return pd.read_csv(path, sep="\t", engine="python", dtype=str, keep_default_na=False)
                except Exception:
                    return pd.read_csv(path, on_bad_lines="skip", engine="python", dtype=str, keep_default_na=False)
    return None


def build_supplementary_docx():
    s1 = read_csv(f"{OUT_DIR}/Supplementary_Table_S1.csv")
    s2 = read_csv(f"{OUT_DIR}/Supplementary_Table_S2.csv")
    s3 = read_csv(f"{OUT_DIR}/Supplementary_Table_S3.csv")
    s4 = read_csv(f"{OUT_DIR}/Supplementary_Table_S4.csv")
    s5 = read_csv(f"{OUT_DIR}/Supplementary_Table_S5.csv")
    s6 = read_csv(f"{OUT_DIR}/Supplementary_Table_S6.csv")
    s7a = read_csv(f"{OUT_DIR}/Supplementary_Table_S7_Process.csv")
    s7b = read_csv(f"{OUT_DIR}/Supplementary_Table_S7_TopModels.csv")
    s8 = read_csv(f"{OUT_DIR}/Supplementary_Table_S8.csv")
    s9 = read_csv(f"{OUT_DIR}/Supplementary_Table_S9.csv")

    if s7a is not None:
        s7a = s7a.copy()
        s7a.insert(0, "Section", "Process")
    if s7b is not None:
        s7b = s7b.copy()
        s7b.insert(0, "Section", "TopModels")
    s7 = None
    if s7a is not None and s7b is not None:
        all_cols = list(dict.fromkeys(list(s7a.columns) + list(s7b.columns)))
        s7a = s7a.reindex(columns=all_cols)
        s7b = s7b.reindex(columns=all_cols)
        s7 = pd.concat([s7a, s7b], ignore_index=True)
    elif s7a is not None:
        s7 = s7a
    elif s7b is not None:
        s7 = s7b
    if s7 is not None:
        s7.to_csv(f"{OUT_DIR}/Supplementary_Table_S7.csv", index=False)

    doc = Document()
    title = doc.add_paragraph("Supplementary Tables")
    style_paragraph(title, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)

    add_table_from_df(doc, "Table S1. Core MR results", s1)
    add_table_from_df(doc, "Table S2. MR sensitivity analyses", s2)
    add_table_from_df(doc, "Table S3. Extended MVMR analyses", s3)
    add_table_from_df(doc, "Table S4. Full colocalization results", s4)
    add_table_from_df(doc, "Table S5. CpG-gene annotation", s5)
    add_table_from_df(doc, "Table S6. Full M12 model coefficients", s6)
    add_table_from_df(doc, "Table S7. Feature-selection process and top models", s7)
    add_table_from_df(doc, "Table S8. Final M12 model performance summary", s8)
    add_table_from_df(doc, "Table S9. Diagnostic metrics (DCA + ROC + calibration)", s9)

    style_doc(doc)
    doc.save(SUPP_DOC)


def strict_format_doc(path):
    doc = Document(path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                if not txt:
                    continue
                new_txt = re.sub(r"(?<=\d)-(?=\d)", "–", txt)
                if new_txt != txt:
                    cell.text = new_txt
    style_doc(doc)
    doc.save(path)


def refine_main_doc_content(path):
    doc = Document(path)
    if len(doc.tables) >= 3:
        t3 = doc.tables[2]
        hdr = [c.text.strip() for c in t3.rows[0].cells]
        try:
            tx_idx = hdr.index("Transcript")
            pr_idx = hdr.index("Protein")
        except ValueError:
            tx_idx, pr_idx = None, None
        if tx_idx is not None and pr_idx is not None:
            mapping = {
                "MFAP2": ("Decrease", "Increase"),
                "CDK11A": ("Not significant", "Increase"),
                "WRAP73": ("Decrease", "Decrease"),
            }
            for i in range(1, len(t3.rows)):
                gene = t3.cell(i, 0).text.strip()
                if gene in mapping:
                    t3.cell(i, tx_idx).text = mapping[gene][0]
                    t3.cell(i, pr_idx).text = mapping[gene][1]
    if len(doc.tables) >= 5:
        t5 = doc.tables[4]
        hdr = [c.text.strip() for c in t5.rows[0].cells]
        hr_idx = next((i for i, h in enumerate(hdr) if "HR (" in h), None)
        cut_idx = next((i for i, h in enumerate(hdr) if "Cutoff method" in h), None)
        ds_idx = next((i for i, h in enumerate(hdr) if h == "Dataset"), 0)
        if hr_idx is not None:
            t5.cell(0, hr_idx).text = "HR (high vs low risk, Cox model, 95% CI)"
        if cut_idx is not None:
            for i in range(1, len(t5.rows)):
                ds = t5.cell(i, ds_idx).text.strip().lower()
                if "tcga" in ds:
                    t5.cell(i, cut_idx).text = "Median cutoff derived from the training cohort"
                else:
                    t5.cell(i, cut_idx).text = "Maxstat-derived optimal cutoff (external cohorts)"
    doc.save(path)


if __name__ == "__main__":
    build_supplementary_docx()
    refine_main_doc_content(MAIN_DOC)
    strict_format_doc(MAIN_DOC)
    strict_format_doc(SUPP_DOC)
    print("done")
