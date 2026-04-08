from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

in_file = "E:/GWAS/Final_Tables_SCI_Format/Main_Tables_Final.docx"
out_file = "E:/GWAS/Final_Tables_SCI_Format/Main_Tables_Final.docx"

doc = Document(in_file)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)

for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = "Times New Roman"
        if r.font.size is None:
            r.font.size = Pt(10.5)
    fmt = p.paragraph_format
    if fmt.space_before is None:
        fmt.space_before = Pt(0)
    if fmt.space_after is None:
        fmt.space_after = Pt(3)
    if p.text.strip().startswith("**Table ") or p.text.strip().startswith("Table "):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(11)

for table in doc.tables:
    if len(table.rows) == 0:
        continue
    header = table.rows[0]
    for cell in header.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
                r.bold = True
    for row in table.rows[1:]:
        for cidx, cell in enumerate(row.cells):
            txt = cell.text.strip().replace(",", "").replace("%", "")
            numeric_like = False
            if txt:
                try:
                    float(txt.split(" ")[0].replace("–", "-"))
                    numeric_like = True
                except Exception:
                    numeric_like = any(ch.isdigit() for ch in txt) and len(txt) < 20
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if numeric_like else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)

doc.save(out_file)
print("formatted")
