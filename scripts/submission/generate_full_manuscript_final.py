import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

out_file = "E:/GWAS/Final_Submission_Figures/Manuscript_Final_Revised.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
for section in doc.sections:
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)
    r.bold = True
    p.paragraph_format.space_after = Pt(10)

def add_heading(text, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

def add_labeled_paragraph(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label)
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(" " + text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

def add_center_paragraph(text, size=11, bold=False, italic=False, space_after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic

def add_formula(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.italic = True

def add_marker(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[Insert {text} here]")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    r.italic = True
    r.bold = True

def add_reference_entry(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

title = "From genetic liability to epigenetic risk: a multi-omics framework for prognostic stratification in lung squamous cell carcinoma"
authors_line = "Xiongjie Li¹†, Fengyue Zhang²†, Zhenyao Wu³, Xuan Xu³, Xiaoyan Zhang⁴, Xianghui Wang¹*"
affil_1 = "¹ Department of Cardiothoracic and Vascular Surgery, Jingzhou Hospital Affiliated to Yangtze University, No. 26 Chuyuan Avenue, Jingzhou District, Jingzhou 434020, Hubei, China"
affil_2 = "² Department of Gastroenterology, Qianjiang Jianghan Oilfield General Hospital, Qianjiang 433100, Hubei, China"
affil_3 = "³ Department of Cardiothoracic, Thyroid and Breast Surgery, Qianjiang Jianghan Oilfield General Hospital, Qianjiang 433100, Hubei, China"
affil_4 = "⁴ Department of Clinical Laboratory, Qianjiang Jianghan Oilfield General Hospital, Qianjiang 433100, Hubei, China"
equal_note = "† These authors contributed equally to this work."
corr_1 = "*Corresponding author: Xianghui Wang (wangxianghui8104@163.com)"
orcid_1 = "Xiongjie Li (ORCID: 0009-0005-9507-6829; Email: surobert518@163.com)"

intro_p1 = "Lung cancer remains a leading cause of cancer-related mortality worldwide [33,36]. Recurrence after curative-intent resection is a major determinant of outcome in non-small cell lung cancer, including LUSC [27,28]. Current clinicopathological risk stratification, including TNM staging, does not fully capture inter-individual variability in recurrence risk [14]. Despite advances in treatment, the identification of patients at high risk of relapse remains limited in LUSC [23], particularly in the context of the heterogeneous molecular landscape of the disease [5,37]."
intro_p2 = "Accumulating evidence suggests that systemic metabolic and inflammatory liabilities contribute to lung cancer development and progression [15,25]. However, such associations are often difficult to interpret because of residual confounding, reverse causation, and population heterogeneity. Mendelian randomization (MR) has emerged as a rigorous framework for strengthening causal inference by leveraging genetic variants as instrumental variables [4,18]. Nevertheless, etiological signals derived from MR alone are insufficient for clinical translation because they do not directly resolve tumor-level molecular mechanisms or prognostic biomarkers."
intro_p3 = "A key missing component is a molecular “translation layer” linking upstream liabilities to stable and clinically measurable tumor features. DNA methylation is well suited to this role because of its biological relevance, relative stability, and mechanistic coupling to cellular metabolism [9,22]. Its translational appeal is further supported by established links between epigenetic remodeling, cancer biology, and biomarker development [20,42]. Partial genetic regulation through methylation quantitative trait loci (mQTL) further enables causal anchoring [11,26]. Integrating MR with mQTL mapping and Bayesian colocalization therefore provides an opportunity to prioritize epigenetic signals that are both genetically anchored and functionally plausible [13,29], while improving the specificity of candidate biomarker discovery through shared-signal evaluation [30,43]."
intro_p4 = "In this study, we propose a causal-to-translational multi-omics framework that integrates Mendelian randomization, mQTL mapping, Bayesian colocalization, and multi-layer omics validation to identify biologically grounded epigenetic signals [17]. We further translate these genetically informed CpG sites into a parsimonious and externally validated methylation-based prognostic model (M12) for risk stratification in lung squamous cell carcinoma. By bridging causal inference with tumor-level molecular features and clinical outcomes, this approach aims to establish a mechanistically informed pathway from genetic liability to prognostic biomarkers in precision oncology."

abstract_background = "Background: Metabolic and inflammatory liabilities have been implicated in lung cancer risk, yet the molecular mechanisms linking upstream exposures to tumor behavior and clinical outcomes remain unclear. Integrative frameworks bridging causal inference and translational biomarkers are needed."
abstract_methods = "Methods: We developed a causal-to-translational multi-omics framework integrating Mendelian randomization (MR), methylation quantitative trait loci (mQTL), Bayesian colocalization, and multi-layer omics validation. Genetically informed CpG sites were prioritized and incorporated into a methylation-based prognostic model using penalized Cox regression. Internal validation was performed in TCGA-LUSC, with external validation in GSE39279 and GSE30219."
abstract_results = "Results: MR analyses revealed a subtype-specific metabolic–inflammatory liability architecture, with stronger effects observed in lung squamous cell carcinoma (LUSC) than lung adenocarcinoma. Two-step MR suggested C-reactive protein (CRP) as a potential mediator linking metabolic traits to LUSC risk. Epigenetic triangulation prioritized MFAP2-associated CpG sites supported by concordant mQTL, MR, and colocalization evidence. Multi-omics analyses converged on a coherent tumor program characterized by platelet activation, coagulation, and extracellular matrix remodeling. A 12-CpG methylation signature (M12) demonstrated stable prognostic performance in the TCGA cohort (C-index = 0.664; 3-year AUC = 0.685), with good calibration (slope = 1.23; R² = 0.84). External validation confirmed significant survival stratification across independent cohorts."
abstract_conclusions = "Conclusions: This study establishes a causal-to-translational framework linking genetic liability to biologically grounded epigenetic markers. The M12 CpG-based signature enables clinically relevant risk stratification in LUSC and provides a scalable strategy for mechanistically informed biomarker development."
keywords = "Keywords: Mendelian randomization; multi-omics integration; DNA methylation; epigenetics; MFAP2; lung squamous cell carcinoma; risk stratification; prognostic model"

methods_21 = "We established a causal-to-translational multi-omics analytical framework to identify biologically informed epigenetic biomarkers and to assess their clinical relevance in lung squamous cell carcinoma (LUSC). The workflow integrated Mendelian randomization (MR) with methylation quantitative trait locus (mQTL) mapping to anchor upstream liabilities to CpG-level regulation [4,13]. Bayesian colocalization and multi-layer omics validation were then used to refine candidate signals and support downstream prognostic application [17,30]. A schematic overview of the study design is provided in Figure 1."
methods_22a = "Two-sample MR analyses were conducted to estimate the putative causal effects of metabolic, inflammatory, and lifestyle-related exposures on lung cancer outcomes. Instrumental variable selection followed established MR guidance, with variants retained at genome-wide significance (P < 5 × 10⁻⁸) and pruned for linkage disequilibrium (r² < 0.001) to ensure independence among retained variants [4]. Summary statistics were obtained from large-scale public GWAS consortia and harmonized through established MR-oriented resources [1,18,39]."
methods_22b = "The primary causal estimate was derived using the inverse-variance weighted (IVW) approach:"
methods_22c = "where w_i denotes the inverse variance of the SNP-outcome association, defined as 1 / (SE_Y,i)²."
methods_22d = "To evaluate the robustness of the IVW estimates and to assess potential horizontal pleiotropy, complementary sensitivity analyses were performed using MR-Egger regression, the weighted median estimator, and MR-PRESSO [3,41]. The MR-Egger model was specified as:"
methods_22e = "Detailed MR sensitivity analyses and supporting results are provided in Tables S1–S2."
methods_23 = "Multivariable MR (MVMR) was performed to estimate the independent effects of correlated exposures within a joint modeling framework [32]. To further explore biologically plausible intermediate pathways, two-step MR analyses were undertaken by separately modeling exposure-mediator and mediator-outcome associations, with particular emphasis on the potential mediating role of inflammatory markers in metabolic risk transmission [6,29]. Indirect effects were interpreted according to the product-of-coefficients framework. Extended MVMR and mediation analyses are shown in Tables S2–S3."
methods_24 = "To prioritize CpG sites that potentially translate inherited liability into tumor-level molecular regulation, we implemented an epigenetic triangulation framework integrating mQTL mapping, MR directionality, and Bayesian colocalization analysis. mQTL resources were first used to identify genetically regulated CpG loci [11,26]. These candidates were then filtered according to concordant MR directionality within a triangulation framework [13,29]. Colocalization analysis was subsequently applied to evaluate whether the exposure-associated and methylation-associated signals were likely attributable to a shared causal variant [30,43]; posterior probability for hypothesis 4 (PP.H4) > 0.30 was considered suggestive evidence of colocalization. CpG sites were retained only when the inferred directions of effect were concordant across the MR and mQTL layers. Full colocalization results are provided in Table S4, with regional plots shown in Figure S3."
methods_25 = "Prioritized CpG sites and their annotated genes were subsequently evaluated across transcriptomic, proteomic, and immune-related datasets to strengthen biological interpretation and translational relevance [17]. Transcriptomic processing and differential analyses were aligned with established RNA-seq and microarray workflows [24,31], with batch correction performed using empirical Bayes adjustment where required [19]. Functional enrichment analyses were performed using established pathway resources, including Reactome and GSVA [10,16], with tissue-level contextualization from GTEx [38] and particular emphasis on tumor immune microenvironment remodeling [2]. Enrichment visualization and annotation were implemented with clusterProfiler [45]. CpG-to-gene mapping details are summarized in Table S5."
methods_26a = "DNA methylation profiles and corresponding clinical follow-up data were obtained from the TCGA-LUSC cohort for prognostic model development [37]. Feature selection was performed using a penalized Cox proportional hazards model with LASSO regularization and 10-fold cross-validation [34,44]. Model development and reporting were aligned with contemporary principles for biomedical prediction modeling [7,35]. Candidate signatures were subsequently compared using a composite selection score, and the final 12-CpG signature (M12_Top12) was selected on the basis of overall parsimony-performance balance. The Cox proportional hazards model was defined as:"
methods_26b = "For each patient, the prognostic risk score was calculated from the multivariable Cox model using z-scored CpG features, where Z(CpG) represents the standardized methylation value of a given CpG site. The general form of the risk score was:"
methods_26c = "The final M12 signature was specified as:"
methods_26d = "Full model coefficients and feature-selection procedures are provided in Tables S6–S7."
methods_27 = "Model performance was comprehensively evaluated using Kaplan-Meier survival analysis, nonparametric time-to-event estimation under censoring, the concordance index (C-index), time-dependent receiver operating characteristic (ROC) analysis, calibration curves, and decision curve analysis (DCA) [40]. In the TCGA-LUSC training cohort, patients were dichotomized into high- and low-risk groups according to the median training-cohort risk score. The TCGA cohort was randomly split into training and internal validation sets using stratified sampling based on survival status. External validation was performed in two independent microarray cohorts (GSE39279 and GSE30219), in which optimal cutoffs were determined using the maxstat (maximally selected rank statistics) procedure. Given the heterogeneity in endpoint definition across validation cohorts, namely recurrence-free survival in GSE39279 and overall survival in GSE30219, external validation primarily focused on discriminative performance, whereas calibration and DCA were principally assessed in the TCGA training cohort. Model development and reporting were aligned with contemporary prediction-model guidance [7,35,44]. Additional diagnostic analyses are shown in Figure S5 and Tables S8–S9."
methods_28 = "All statistical analyses were performed using R software (version 4.5.2). Unless otherwise specified, all statistical tests were two-sided, and P-values < 0.05 were considered statistically significant. Where appropriate, results were interpreted in conjunction with multiple-testing considerations and consistency across analytical layers."

res_31a = "We first assessed the causal effects of metabolic, inflammatory, and lifestyle exposures on lung cancer outcomes using Mendelian randomization. Smoking initiation exhibited the strongest and most consistent associations across outcomes, with larger effect sizes observed in lung squamous cell carcinoma (LUSC) compared with lung adenocarcinoma (LUAD). Body mass index (BMI) showed significant associations with LUSC but not LUAD, suggesting subtype-specific susceptibility patterns (Figure 2; Table 2)."
res_31b = "Detailed effect estimates, including odds ratios and confidence intervals, are provided in Table 2."
res_31c = "Multivariable MR analyses attenuated these associations, consistent with shared genetic architecture among cardiometabolic traits. Collectively, these findings support a networked metabolic–inflammatory liability framework rather than a single dominant causal factor (Figure S2; Tables S1–S2)."
res_32a = "To explore potential mechanisms linking metabolic traits to LUSC, we performed two-step Mendelian randomization analyses. In particular, C-reactive protein (CRP) emerged as a potential mediator, with the BMI→CRP→LUSC pathway showing consistent exposure–mediator and mediator–outcome effects."
res_32b = "Although these associations did not survive multiple testing correction, the recurrent involvement of CRP across multiple pathways suggests a biologically plausible inflammatory axis underlying LUSC risk, consistent with prior evidence linking systemic inflammation to lung carcinogenesis [15] (Figure 2C; Tables S2–S3)."
res_33a = "To bridge genetic liability to tumor-level regulation, we implemented an epigenetic triangulation framework integrating mQTL mapping, Mendelian randomization, and colocalization analysis (see Methods). This approach enabled the identification of CpG-specific regulatory signals."
res_33b = "At the MFAP2 locus, one CpG site demonstrated concordant evidence across all layers, whereas adjacent probes did not, indicating probe-specific functional relevance. Colocalization analysis further supported a shared genetic signal at this locus."
res_33c = "These findings prioritized MFAP2 as an upstream candidate node linking inherited liabilities to epigenetic regulation and downstream tumor biology (Figure 3; Table 3; Figure S3; Table S4)."
res_34a = "To further characterize the biological relevance of prioritized loci, we performed multi-omics integration across transcriptomic, proteomic, and immune-related datasets. These analyses converged on a coherent tumor program characterized by platelet activation, coagulation processes, and extracellular matrix remodeling, with Figure 6 further illustrating CpG–pathway associations and immune-related signatures (Figures 4 and 6)."
res_34b = "Notably, transcriptomic and proteomic analyses showed partial discordance at the gene expression level but consistent activation at the pathway level, suggesting post-transcriptional regulation. Functional enrichment further supported a platelet–ECM axis and microtubule-related transport processes, in line with prior evidence implicating platelet activation, extracellular matrix remodeling, and metastatic support functions in tumor progression [12,21]."
res_34c = "Together, these results provide a mechanistic bridge linking systemic metabolic–inflammatory liabilities to tumor microenvironment remodeling (Table S5)."
res_35a = "We translated the prioritized epigenetic signals into a methylation-based prognostic model using penalized LASSO-Cox regression. The final model comprised 12 CpG sites (M12) and demonstrated stable prognostic performance in the TCGA-LUSC training cohort, consistent with contemporary principles for parsimonious prognostic modeling [35] (Figure 5A–C; Table 4; Tables S6–S7)."
res_35b = "Using the median training-cohort cutoff, Kaplan–Meier analysis showed significantly worse overall survival in the high-risk group (Figure 5D). Continuous risk modeling further supported the prognostic relevance of the M12 signature. Model coefficients and feature contributions are summarized in Table 4."
res_36a = "We next evaluated model performance across the TCGA internal validation cohort and two independent external cohorts (GSE39279 and GSE30219). Consistent survival stratification was observed in the internal validation subset and across external datasets using maxstat-derived optimal cutoffs, supporting the portability of the M12 model (Figure 7B–D)."
res_36b = "In the TCGA training cohort, the model achieved a C-index of 0.664 and a 3-year AUC of 0.685, with good calibration performance (slope = 1.227; R² = 0.843). Detailed cohort-level performance metrics are summarized in Table 5 and Table S8, whereas calibration analysis demonstrated good agreement between predicted and observed survival in the training cohort (Figure S5A)."
res_36c = "Taken together, these findings indicate that the M12 signature captures clinically meaningful prognostic information and supports both group-based stratification and continuous risk assessment; additional diagnostic analyses further supported model robustness (Figure S5; Table S9)."

disc_1 = "In this study, we developed a causal-to-translational multi-omics framework that links genetic liability to clinically relevant epigenetic markers in lung squamous cell carcinoma (LUSC). By integrating Mendelian randomization with epigenetic triangulation and multi-omics interpretation, our approach moves beyond conventional association-based biomarker discovery and aligns with broader efforts to connect causal inference with molecular profiling in complex disease biology [17]. These signals were subsequently translated into a 12-CpG methylation signature (M12), which demonstrated consistent prognostic stratification in both training and independent validation cohorts, with acceptable discrimination and good calibration. Together, these findings support a unified framework in which upstream metabolic–inflammatory liabilities are connected to tumor microenvironment remodeling and ultimately to clinically actionable risk prediction."
disc_2 = "A key insight from this work is that metabolic–inflammatory risk in LUSC is best conceptualized as a networked liability rather than a single-factor effect [15,25]. The recurrent involvement of CRP across multiple pathways supports an inflammatory axis linking systemic metabolic perturbations to tumor risk. Although these findings should be interpreted as hypothesis-generating rather than definitive mediation, they provide a coherent biological framework for understanding subtype-specific susceptibility."
disc_3 = "The epigenetic triangulation strategy represents an important refinement step by integrating genetic and molecular evidence at the CpG level. The probe-specific signal observed at the MFAP2 locus highlights the importance of fine-grained epigenetic resolution, as adjacent CpG sites within the same locus may exhibit distinct functional roles. This finding underscores the value of combining mQTL, MR, and colocalization to improve the specificity of biomarker discovery [11,30]."
disc_4 = "Multi-omics integration further demonstrated that these genetically anchored signals converge on a coherent tumor program characterized by coagulation, platelet activation, extracellular matrix remodeling, and immune-contexture shifts within the tumor microenvironment [2]. This program provides a plausible mechanistic bridge linking systemic liabilities to tumor microenvironment dynamics, consistent with foundational links between inflammatory microenvironments and tumor promotion [8]. It is also biologically compatible with prior evidence implicating platelet-mediated interactions in metastatic progression [12,21]."
disc_5 = "From a clinical perspective, the M12 methylation signature illustrates that biologically informed multi-feature models can capture meaningful prognostic information. Although its discriminative performance was moderate, calibration was good in the training cohort and significant survival stratification was preserved across independent datasets. The stronger performance of continuous risk assessment further suggests that the signature may be most informative when interpreted as a quantitative prognostic measure, consistent with established recommendations for clinically usable prediction models and epigenetic biomarker translation [35,42]."
disc_6 = "Importantly, beyond the specific findings, this study provides a generalizable framework for biomarker discovery that integrates causal inference with multi-omics validation. This design enables a systematic transition from genetic association to functional prioritization and clinical translation, in line with broader efforts to integrate molecular profiling across complex diseases [17], and may be applicable to other complex diseases."
disc_7 = "Several limitations should be acknowledged. First, MR is designed for etiological inference rather than prognostic modeling, and the transition from genetic liability to recurrence outcomes should be interpreted with caution [6]. Second, colocalization analyses may be affected by locus complexity and the presence of multiple causal variants [43]. Third, external validation was conducted in relatively modest cohorts, which may limit generalizability. Future studies should focus on prospective validation and experimental characterization of prioritized CpG sites."
disc_8 = "In conclusion, our findings support a causal-to-translational paradigm for biomarker discovery, demonstrating how genetically informed epigenetic signals can be leveraged to bridge upstream risk factors and clinical outcomes in LUSC, thereby providing a generalizable blueprint for translating causal inference into clinically actionable biomarkers in complex diseases."

concl_1 = "This study establishes a causal-to-translational paradigm linking genetic liability to tumor biology and clinical outcomes in lung squamous cell carcinoma. Through integrative analyses, we identify MFAP2-linked epigenetic regulation as a key node connecting systemic metabolic–inflammatory risk to tumor progression. The derived M12 CpG-based signature demonstrates that biologically informed models can capture clinically relevant prognostic signals, particularly when interpreted as continuous measures. Beyond the specific findings, this work provides a generalizable framework for translating causal inference into mechanistically grounded biomarkers for precision oncology."
ack_1 = "We acknowledge the investigators and participants of the International Lung Cancer Consortium (ILCCO), The Cancer Genome Atlas (TCGA), and the Clinical Proteomic Tumor Analysis Consortium (CPTAC) for generating and making these data publicly available. We also thank the developers of the analytical tools and open-source resources that made this study possible."
author_1 = "Xiongjie Li (XL): Conceived the study, designed the methodological framework, performed data acquisition and curation, conducted Mendelian randomization analyses and prognostic modeling, and drafted the original manuscript."
author_2 = "Fengyue Zhang (FZ): Participated in data acquisition and curation, Mendelian randomization analyses, and prognostic modeling."
author_3 = "Xuan Xu (XX): Participated in multi-omics data integration and data visualization."
author_4 = "Zhenyao Wu (ZW): Participated in multi-omics data integration and data visualization."
author_5 = "Xiaoyan Zhang (XZ): Provided clinical interpretation and assisted with the external validation analyses."
author_6 = "Xianghui Wang (XW): Supervised the study, provided administrative support, and critically revised the manuscript."
author_7 = "All authors read and approved the final manuscript."
author_8 = "Xiongjie Li and Fengyue Zhang contributed equally to this work."
funding_1 = "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors."
competing_1 = "The authors declare no competing interests."
ethics_1 = "Ethical review and approval were waived for this study because all analyses were based on publicly available, de-identified summary statistics and open-access molecular datasets (e.g., GWAS, TCGA, CPTAC). All original studies obtained appropriate institutional ethical approvals and informed consent from participants, in accordance with the Declaration of Helsinki."
data_1 = "Publicly available datasets were analyzed in this study. GWAS summary statistics were obtained from the IEU OpenGWAS platform (https://gwas.mrcieu.ac.uk/). Transcriptomic and methylation data were accessed through The Cancer Genome Atlas data portal (https://portal.gdc.cancer.gov/). Proteomic data were obtained from the Clinical Proteomic Tumor Analysis Consortium data portal. Code for Mendelian randomization analyses, prognostic modeling, and visualization is available at: https://github.com/morningLxj/Metabolic-LungCancer-MR. All data sources used in this study are publicly accessible, and detailed accession information is provided in the Methods section."
references = [
    "1. Potter, A. L. et al. Recurrence After Complete Resection for Non-Small Cell Lung Cancer in the National Lung Screening Trial. Ann Thorac Surg 116, 684–692, doi:10.1016/j.athoracsur.2023.06.004 (2023).",
    "2. Cancer Genome Atlas Research, N. Comprehensive genomic characterization of squamous cell lung cancers. Nature 489, 519–525, doi:10.1038/nature11404 (2012).",
    "3. Lau, S. C. M., Pan, Y., Velcheti, V. & Wong, K. K. Squamous cell lung cancer: Current landscape and future therapeutic options. Cancer Cell 40, 1279–1293, doi:10.1016/j.ccell.2022.09.018 (2022).",
    "4. Rajaram, R. et al. Recurrence-Free Survival in Patients With Surgically Resected Non-Small Cell Lung Cancer: A Systematic Literature Review and Meta-Analysis. Chest 165, 1260–1270, doi:10.1016/j.chest.2023.11.042 (2024).",
    "5. Carreras-Torres, R. et al. Obesity, metabolic factors and risk of different histological types of lung cancer: A Mendelian randomization study. PLoS One 12, e0177875, doi:10.1371/journal.pone.0177875 (2017).",
    "6. Wei, Z. et al. The association between metabolic syndrome and lung cancer risk: a Mendelian randomization study. Sci Rep 14, 28494, doi:10.1038/s41598-024-79260-y (2024).",
    "7. Zhou, B., Liu, J., Wang, Z. M. & Xi, T. C-reactive protein, interleukin 6 and lung cancer risk: a meta-analysis. PLoS One 7, e43075, doi:10.1371/journal.pone.0043075 (2012).",
    "8. Di Lorenzo, C., Dell'agli, M., Colombo, E., Sangiovanni, E. & Restani, P. Metabolic syndrome and inflammation: a critical review of in vitro and clinical approaches for benefit assessment of plant food supplements. Evid Based Complement Alternat Med 2013, 782461, doi:10.1155/2013/782461 (2013).",
    "9. Sanderson, E. et al. Mendelian randomization. Nat Rev Methods Primers 2, doi:10.1038/s43586-021-00092-5 (2022).",
    "10. Kaelin, W. G., Jr. & McKnight, S. L. Influence of metabolism on epigenetics and disease. Cell 153, 56–69, doi:10.1016/j.cell.2013.03.004 (2013).",
    "11. Burgess, S. et al. Guidelines for performing Mendelian randomization investigations: update for summer 2023. Wellcome Open Res 4, 186, doi:10.12688/wellcomeopenres.15555.3 (2019).",
    "12. Lu, C. & Thompson, C. B. Metabolic regulation of epigenetics. Cell Metab 16, 9–17, doi:10.1016/j.cmet.2012.06.001 (2012).",
    "13. consortium, B. Quantitative comparison of DNA methylation assays for biomarker development and clinical applications. Nat Biotechnol 34, 726–737, doi:10.1038/nbt.3605 (2016).",
    "14. Wagner, W. How to Translate DNA Methylation Biomarkers Into Clinical Practice. Front Cell Dev Biol 10, 854797, doi:10.3389/fcell.2022.854797 (2022).",
    "15. Gaunt, T. R. et al. Systematic identification of genetic influences on methylation across the human life course. Genome Biol 17, 61, doi:10.1186/s13059-016-0926-z (2016).",
    "16. Relton, C. L. & Davey Smith, G. Two-step epigenetic Mendelian randomization: a strategy for establishing the causal role of epigenetic processes in pathways to disease. Int J Epidemiol 41, 161–176, doi:10.1093/ije/dyr233 (2012).",
    "17. Richardson, T. G. et al. Systematic Mendelian randomization framework elucidates hundreds of CpG sites which may mediate the influence of genetic variants on disease. Hum Mol Genet 27, 3293–3304, doi:10.1093/hmg/ddy210 (2018).",
    "18. Locke, A. E. et al. Genetic studies of body mass index yield new insights for obesity biology. Nature 518, 197–206, doi:10.1038/nature14177 (2015).",
    "19. Liu, M. et al. Association studies of up to 1.2 million individuals yield new insights into the genetic etiology of tobacco and alcohol use. Nat Genet 51, 237–244, doi:10.1038/s41588-018-0307-5 (2019).",
    "20. Ligthart, S. et al. Genome Analyses of >200,000 Individuals Identify 58 Loci for Chronic Inflammation and Highlight Pathways that Link Inflammation and Complex Disorders. Am J Hum Genet 103, 691–706, doi:10.1016/j.ajhg.2018.09.009 (2018).",
    "21. Willer, C. J. et al. Discovery and refinement of loci associated with lipid levels. Nat Genet 45, 1274–1283, doi:10.1038/ng.2797 (2013).",
    "22. McKay, J. D. et al. Large-scale association analysis identifies new lung cancer susceptibility loci and heterogeneity in genetic susceptibility across histological subtypes. Nat Genet 49, 1126–1132, doi:10.1038/ng.3892 (2017).",
    "23. Fanidi, A. et al. Circulating Folate, Vitamin B6, and Methionine in Relation to Lung Cancer Risk in the Lung Cancer Cohort Consortium (LC3). J Natl Cancer Inst 110, 57–67, doi:10.1093/jnci/djx119 (2018).",
    "24. Sung, H. et al. Global Cancer Statistics 2020: GLOBOCAN Estimates of Incidence and Mortality Worldwide for 36 Cancers in 185 Countries. CA Cancer J Clin 71, 209–249, doi:10.3322/caac.21660 (2021).",
    "25. Davies, N. M., Holmes, M. V. & Davey Smith, G. Reading Mendelian randomisation studies: a guide, glossary, and checklist for clinicians. BMJ 362, k601, doi:10.1136/bmj.k601 (2018).",
    "26. Hemani, G. et al. The MR-Base platform supports systematic causal inference across the human phenome. Elife 7, doi:10.7554/eLife.34408 (2018).",
    "27. Verbanck, M., Chen, C. Y., Neale, B. & Do, R. Publisher Correction: Detection of widespread horizontal pleiotropy in causal relationships inferred from Mendelian randomization between complex traits and diseases. Nat Genet 50, 1196, doi:10.1038/s41588-018-0164-2 (2018).",
    "28. Consortium, G. T. The GTEx Consortium atlas of genetic regulatory effects across human tissues. Science 369, 1318–1330, doi:10.1126/science.aaz1776 (2020).",
    "29. Giambartolomei, C. et al. Bayesian test for colocalisation between pairs of genetic association studies using summary statistics. PLoS Genet 10, e1004383, doi:10.1371/journal.pgen.1004383 (2014).",
    "30. Love, M. I., Huber, W. & Anders, S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biol 15, 550, doi:10.1186/s13059-014-0550-8 (2014).",
    "31. Robinson, M. D., McCarthy, D. J. & Smyth, G. K. edgeR: a Bioconductor package for differential expression analysis of digital gene expression data. Bioinformatics 26, 139–140, doi:10.1093/bioinformatics/btp616 (2010).",
    "32. Tanaka, T. [[Fundamentals] 5. Python+scikit-learn for Machine Learning in Medical Imaging]. Nihon Hoshasen Gijutsu Gakkai Zasshi 79, 1189–1193, doi:10.6009/jjrt.2023-2266 (2023).",
    "33. Troyanskaya, O. et al. Missing value estimation methods for DNA microarrays. Bioinformatics 17, 520–525, doi:10.1093/bioinformatics/17.6.520 (2001).",
    "34. Johnson, W. E., Li, C. & Rabinovic, A. Adjusting batch effects in microarray expression data using empirical Bayes methods. Biostatistics 8, 118–127, doi:10.1093/biostatistics/kxj037 (2007).",
    "35. Kanehisa, M. & Goto, S. KEGG: kyoto encyclopedia of genes and genomes. Nucleic Acids Res 28, 27–30, doi:10.1093/nar/28.1.27 (2000).",
    "36. Yu, G., Wang, L. G., Han, Y. & He, Q. Y. clusterProfiler: an R package for comparing biological themes among gene clusters. OMICS 16, 284–287, doi:10.1089/omi.2011.0118 (2012).",
    "37. Fabregat, A. et al. The Reactome pathway Knowledgebase. Nucleic Acids Res 44, D481–487, doi:10.1093/nar/gkv1351 (2016).",
    "38. Hanzelmann, S., Castelo, R. & Guinney, J. GSVA: gene set variation analysis for microarray and RNA-seq data. BMC Bioinformatics 14, 7, doi:10.1186/1471-2105-14-7 (2013).",
    "39. Friedman, J., Hastie, T. & Tibshirani, R. Regularization Paths for Generalized Linear Models via Coordinate Descent. J Stat Softw 33, 1–22 (2010).",
    "40. Dinse, G. E. & Lagakos, S. W. Nonparametric estimation of lifetime and disease onset distributions from incomplete observations. Biometrics 38, 921–932 (1982).",
    "41. Goldstraw, P. et al. The IASLC Lung Cancer Staging Project: Proposals for Revision of the TNM Stage Groupings in the Forthcoming (Eighth) Edition of the TNM Classification for Lung Cancer. J Thorac Oncol 11, 39–51, doi:10.1016/j.jtho.2015.09.009 (2016).",
    "42. Satpathy, S. et al. A proteogenomic portrait of lung squamous cell carcinoma. Cell 184, 4348–4371 e4340, doi:10.1016/j.cell.2021.07.016 (2021).",
    "43. Bowden, J., Davey Smith, G. & Burgess, S. Mendelian randomization with invalid instruments: effect estimation and bias detection through Egger regression. Int J Epidemiol 44, 512–525, doi:10.1093/ije/dyv080 (2015).",
    "44. Smith, G. D. & Ebrahim, S. 'Mendelian randomization': can genetic epidemiology contribute to understanding environmental determinants of disease? Int J Epidemiol 32, 1–22, doi:10.1093/ije/dyg070 (2003).",
    "45. Wallace, C. A more accurate method for colocalisation analysis allowing for multiple causal variants. PLoS Genet 17, e1009440, doi:10.1371/journal.pgen.1009440 (2021).",
    "46. Min, J. L. et al. Genomic and phenotypic insights from an atlas of genetic effects on DNA methylation. Nat Genet 53, 1311–1321, doi:10.1038/s41588-021-00923-x (2021).",
    "47. Sanderson, E., Davey Smith, G., Windmeijer, F. & Bowden, J. An examination of multivariable Mendelian randomization in the single-sample and two-sample summary data settings. Int J Epidemiol 48, 713–727, doi:10.1093/ije/dyy262 (2019).",
    "48. Mantovani, A., Allavena, P., Sica, A. & Balkwill, F. Cancer-related inflammation. Nature 454, 436–444, doi:10.1038/nature07205 (2008).",
    "49. Coussens, L. M. & Werb, Z. Inflammation and cancer. Nature 420, 860–867, doi:10.1038/nature01322 (2002).",
    "50. Carter, A. R. et al. Mendelian randomisation for mediation analysis: current methods and challenges for implementation. Eur J Epidemiol 36, 465–478, doi:10.1007/s10654-021-00757-1 (2021).",
    "51. Gay, L. J. & Felding-Habermann, B. Contribution of platelets to tumour metastasis. Nat Rev Cancer 11, 123–134, doi:10.1038/nrc3004 (2011).",
    "52. Laird, P. W. The power and the promise of DNA methylation markers. Nat Rev Cancer 3, 253–266, doi:10.1038/nrc1045 (2003)."
]

ref_list_file = "E:/GWAS/Final_Submission_Figures/参考文献列表.txt"
with open(ref_list_file, "r", encoding="utf-8") as f:
    references = [line.strip() for line in f if re.match(r"^\d+\.\s", line.strip())]

add_title(title)
add_center_paragraph(authors_line, size=11, bold=False, italic=False, space_after=3)
add_center_paragraph(affil_1, size=10.5, space_after=1)
add_center_paragraph(affil_2, size=10.5, space_after=1)
add_center_paragraph(affil_3, size=10.5, space_after=1)
add_center_paragraph(affil_4, size=10.5, space_after=2)
add_center_paragraph(equal_note, size=10.5, italic=True, space_after=1)
add_center_paragraph(corr_1, size=10.5, space_after=1)
add_center_paragraph(orcid_1, size=10.5, space_after=6)
add_heading("Abstract", size=13)
add_labeled_paragraph("Background:", abstract_background.replace("Background: ", ""))
add_labeled_paragraph("Methods:", abstract_methods.replace("Methods: ", ""))
add_labeled_paragraph("Results:", abstract_results.replace("Results: ", ""))
add_labeled_paragraph("Conclusions:", abstract_conclusions.replace("Conclusions: ", ""))
add_labeled_paragraph("Keywords:", keywords.replace("Keywords: ", ""))

add_heading("1. Introduction", size=13)
add_paragraph(intro_p1)
add_paragraph(intro_p2)
add_paragraph(intro_p3)
add_paragraph(intro_p4)

add_heading("2. Methods", size=13)
add_heading("2.1 Study design and analytical framework")
add_paragraph(methods_21)
add_marker("Figure 1")
add_heading("2.2 Genetic liability analysis using Mendelian randomization")
add_paragraph(methods_22a)
add_paragraph(methods_22b)
add_formula("β̂_IVW = ∑ [ w_i · (β_Y,i / β_X,i) ] / ∑ w_i")
add_paragraph(methods_22c)
add_paragraph(methods_22d)
add_formula("β_Y,i = α + β_MR · β_X,i + ε_i")
add_paragraph(methods_22e)
add_heading("2.3 Multivariable and mediation MR")
add_paragraph(methods_23)
add_heading("2.4 Epigenetic triangulation framework")
add_paragraph(methods_24)
add_heading("2.5 Multi-omics integration and functional annotation")
add_paragraph(methods_25)
add_heading("2.6 Prognostic model construction (M12 Signature)")
add_paragraph(methods_26a)
add_formula("h(t|X) = h_0(t) exp(β_1X_1 + ... + β_pX_p)")
add_paragraph(methods_26b)
add_formula("Risk Score = ∑ (β_j × Z(CpG_j))")
add_paragraph(methods_26c)
add_formula("Risk Score = (0.210 × Z[cg06238570]) + (0.140 × Z[cg07318658]) + (0.104 × Z[cg13144823]) - (0.222 × Z[cg14526939]) - (0.294 × Z[cg07151966]) + (0.041 × Z[cg12672785]) + (0.075 × Z[cg05801080]) - (0.205 × Z[cg07154958]) + (0.253 × Z[cg16098170]) + (0.260 × Z[cg09010499]) - (0.230 × Z[cg05984948]) - (0.152 × Z[cg00646216])")
add_paragraph(methods_26d)
add_heading("2.7 Model evaluation and validation")
add_paragraph(methods_27)
add_heading("2.8 Statistical analysis")
add_paragraph(methods_28)

add_heading("3. Results", size=13)
add_heading("3.1 Genetic liability reveals subtype-specific risk architecture")
add_paragraph(res_31a)
add_paragraph(res_31b)
add_paragraph(res_31c)
add_marker("Figure 2 and Table 2")
add_heading("3.2 Inflammatory pathways may mediate metabolic risk in LUSC")
add_paragraph(res_32a)
add_paragraph(res_32b)
add_heading("3.3 Epigenetic triangulation prioritizes MFAP2-associated CpG sites")
add_paragraph(res_33a)
add_paragraph(res_33b)
add_paragraph(res_33c)
add_marker("Figure 3 and Table 3")
add_heading("3.4 Multi-omics integration reveals a coagulation–ECM tumor program")
add_paragraph(res_34a)
add_paragraph(res_34b)
add_paragraph(res_34c)
add_marker("Figures 4 and 6")
add_heading("3.5 A 12-CpG methylation signature enables clinically relevant risk stratification")
add_paragraph(res_35a)
add_marker("Figure 5A–C and Table 4")
add_paragraph(res_35b)
add_marker("Figure 5D")
add_heading("3.6 External validation confirms model portability across independent cohorts")
add_paragraph(res_36a)
add_paragraph(res_36b)
add_marker("Figure 7 and Table 5")
add_paragraph(res_36c)

add_heading("4. Discussion", size=13)
add_paragraph(disc_1)
add_paragraph(disc_2)
add_paragraph(disc_3)
add_paragraph(disc_4)
add_paragraph(disc_5)
add_paragraph(disc_6)
add_paragraph(disc_7)
add_paragraph(disc_8)

add_heading("5. Conclusion", size=13)
add_paragraph(concl_1)

add_heading("Acknowledgements", size=13)
add_paragraph(ack_1)

add_heading("Author contributions", size=13)
add_paragraph(author_1)
add_paragraph(author_2)
add_paragraph(author_3)
add_paragraph(author_4)
add_paragraph(author_5)
add_paragraph(author_6)
add_paragraph(author_7)
add_paragraph(author_8)

add_heading("Funding", size=13)
add_paragraph(funding_1)

add_heading("Declarations", size=13)
add_heading("Competing interests")
add_paragraph(competing_1)
add_heading("Ethics approval and consent to participate")
add_paragraph(ethics_1)
add_heading("Availability of data and materials")
add_paragraph(data_1)

add_heading("References", size=13)
for ref in references:
    add_reference_entry(ref)

saved_path = out_file
try:
    doc.save(out_file)
except PermissionError:
    saved_path = out_file.replace(".docx", "_latest.docx")
    doc.save(saved_path)
print(f"Saved to {saved_path}")
