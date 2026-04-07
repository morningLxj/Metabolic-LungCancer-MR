import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

out_file = "E:/GWAS/Final_Submission_Figures/Methods_Revised.docx"

doc = Document()

# Define styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)

def add_heading(text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14 if level == 1 else 12)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)

add_heading("2. Methods", level=1)

add_heading("2.1 Study design and analytical framework")
add_paragraph("We established a causal-to-translational multi-omics analytical framework to identify biologically informed epigenetic biomarkers and to assess their clinical relevance in lung squamous cell carcinoma (LUSC). The workflow integrated Mendelian randomization (MR), methylation quantitative trait locus (mQTL) mapping, Bayesian colocalization, and multi-layer omics validation, thereby enabling stepwise prioritization from upstream genetic liability to downstream prognostic application. A schematic overview of the study design is provided in Figure 1.")

add_heading("2.2 Genetic liability analysis using Mendelian randomization")
add_paragraph("Two-sample MR analyses were conducted to estimate the putative causal effects of metabolic, inflammatory, and lifestyle-related exposures on lung cancer outcomes. Instrumental variables were selected at genome-wide significance (P < 5 × 10⁻⁸) and pruned for linkage disequilibrium (r² < 0.001) to ensure independence among retained variants.")
add_paragraph("The primary causal estimate was derived using the inverse-variance weighted (IVW) approach:")

ivw_formula = "β̂_IVW = ∑ [ w_i · (β_Y,i / β_X,i) ] / ∑ w_i"
p_ivw = doc.add_paragraph(ivw_formula)
p_ivw.paragraph_format.left_indent = Pt(24)
p_ivw.runs[0].font.name = 'Times New Roman'
p_ivw.runs[0].italic = True

add_paragraph("where w_i denotes the inverse variance of the SNP-outcome association, defined as 1 / (SE_Y,i)².")

add_paragraph("To evaluate the robustness of the IVW estimates and to assess potential horizontal pleiotropy, complementary sensitivity analyses were performed using MR-Egger regression, the weighted median estimator, and MR-PRESSO. The MR-Egger model was specified as:")

egger_formula = "β_Y,i = α + β_MR · β_X,i + ε_i"
p_eg = doc.add_paragraph(egger_formula)
p_eg.paragraph_format.left_indent = Pt(24)
p_eg.runs[0].font.name = 'Times New Roman'
p_eg.runs[0].italic = True
doc.add_paragraph("")

add_heading("2.3 Multivariable and mediation MR")
add_paragraph("Multivariable MR (MVMR) was performed to estimate the independent effects of correlated exposures within a joint modeling framework. To further explore biologically plausible intermediate pathways, two-step MR analyses were undertaken by separately modeling exposure-mediator and mediator-outcome associations, with particular emphasis on the potential mediating role of inflammatory markers in metabolic risk transmission. Indirect effects were interpreted according to the product-of-coefficients framework.")

add_heading("2.4 Epigenetic triangulation framework")
add_paragraph("To prioritize CpG sites that potentially translate inherited liability into tumor-level molecular regulation, we implemented an epigenetic triangulation framework integrating mQTL mapping, MR directionality, and Bayesian colocalization analysis. mQTL datasets were first used to identify genetically regulated CpG loci. Colocalization analysis was then applied to evaluate whether the exposure-associated and methylation-associated signals were likely attributable to a shared causal variant; posterior probability for hypothesis 4 (PP.H4) > 0.30 was considered suggestive evidence of colocalization. CpG sites were retained only when the inferred directions of effect were concordant across the MR and mQTL layers.")

add_heading("2.5 Multi-omics integration and functional annotation")
add_paragraph("Prioritized CpG sites and their annotated genes were subsequently evaluated across transcriptomic, proteomic, and immune-related datasets to strengthen biological interpretation and translational relevance. Functional enrichment analyses were performed to identify pathways associated with the candidate genes, with particular emphasis on coagulation, platelet activation, extracellular matrix organization, and immune-related processes implicated in tumor microenvironment remodeling.")

add_heading("2.6 Prognostic model construction (M12 Signature)")
add_paragraph("DNA methylation profiles and corresponding clinical follow-up data were obtained from the TCGA-LUSC cohort for prognostic model development. Feature selection was performed using a penalized Cox proportional hazards model with LASSO regularization and 10-fold cross-validation. Candidate signatures were subsequently compared using a composite selection score, and the final 12-CpG signature (M12_Top12) was selected on the basis of overall parsimony-performance balance. The Cox proportional hazards model was defined as:")

cox_formula = "h(t|X) = h_0(t) exp(β_1X_1 + ... + β_pX_p)"
p_cox = doc.add_paragraph(cox_formula)
p_cox.paragraph_format.left_indent = Pt(24)
p_cox.runs[0].font.name = 'Times New Roman'
p_cox.runs[0].italic = True

add_paragraph("For each patient, the prognostic risk score was calculated from the multivariable Cox model using z-scored CpG features, where Z(CpG) represents the standardized methylation value of a given CpG site. The general form of the risk score was:")

general_formula = "Risk Score = ∑ (β_j × Z(CpG_j))"
p_general = doc.add_paragraph(general_formula)
p_general.paragraph_format.left_indent = Pt(24)
p_general.runs[0].font.name = 'Times New Roman'
p_general.runs[0].italic = True

add_paragraph("The final M12 signature was specified as:")

formula_text = (
    "Risk Score = (0.210 × Z[cg06238570]) + (0.140 × Z[cg07318658]) + (0.104 × Z[cg13144823]) "
    "- (0.222 × Z[cg14526939]) - (0.294 × Z[cg07151966]) + (0.041 × Z[cg12672785]) "
    "+ (0.075 × Z[cg05801080]) - (0.205 × Z[cg07154958]) + (0.253 × Z[cg16098170]) "
    "+ (0.260 × Z[cg09010499]) - (0.230 × Z[cg05984948]) - (0.152 × Z[cg00646216])"
)
p_form = doc.add_paragraph(formula_text)
p_form.paragraph_format.left_indent = Pt(24)
p_form.paragraph_format.line_spacing = 1.5
for r in p_form.runs:
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.italic = True
doc.add_paragraph("")


add_heading("2.7 Model evaluation and validation")
add_paragraph("Model performance was comprehensively evaluated using Kaplan-Meier survival analysis, the concordance index (C-index), time-dependent receiver operating characteristic (ROC) analysis, calibration curves, and decision curve analysis (DCA). In the TCGA-LUSC training cohort, patients were dichotomized into high- and low-risk groups according to the median training-cohort risk score. External validation was performed in two independent microarray cohorts (GSE39279 and GSE30219), in which optimal cutoffs were determined using the maxstat (maximally selected rank statistics) procedure. Given the heterogeneity in endpoint definition across validation cohorts, namely recurrence-free survival in GSE39279 and overall survival in GSE30219, external validation primarily focused on discriminative performance, whereas calibration and DCA were principally assessed in the TCGA training cohort.")

add_heading("2.8 Statistical analysis")
add_paragraph("All statistical analyses were performed using R software (version 4.5.2). Unless otherwise specified, all statistical tests were two-sided, and P-values < 0.05 were considered statistically significant. Where appropriate, results were interpreted in conjunction with multiple-testing considerations and consistency across analytical layers.")

doc.save(out_file)
print(f"Saved to {out_file}")
