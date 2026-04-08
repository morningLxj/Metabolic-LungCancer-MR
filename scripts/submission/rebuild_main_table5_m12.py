import os
import re
import numpy as np
import pandas as pd
from lifelines import CoxPHFitter
from docx import Document

base = "E:/GWAS"
out_dir = f"{base}/Final_Tables_SCI_Format"
fig7_file = f"{base}/Final_Submission_Figures/Main_Figures_14CpG/Figure7_M12_Summary.csv"
risk_file = f"{base}/PDC/analysis/model_selection_14cpg/BestModel_RiskScores.csv"
g39279_file = f"{base}/GSE39279_LUSC_Correct_Risk_Score_Data.csv"
g30219_file = f"{base}/Final_Figures/GSE30219_parsed_local.csv"
md_file = f"{out_dir}/Main_Tables_Final_current.md"
md_file2 = f"{out_dir}/Main_Tables_Final.updated.md"
docx_file = f"{out_dir}/Main_Tables_Final.docx"

sup_digits = str.maketrans("-0123456789", "⁻⁰¹²³⁴⁵⁶⁷⁸⁹")

def fmt_p(p):
    try:
        v = float(p)
    except Exception:
        return "Not available"
    if np.isnan(v):
        return "Not available"
    prefix = "P = "
    if v < 1e-3:
        e = int(np.floor(np.log10(v)))
        m = v / (10 ** e)
        return f"{prefix}{m:.2f} × 10{str(e).translate(sup_digits)}"
    return f"{prefix}{v:.3f}".rstrip("0").rstrip(".")

def cox_hr_ci_p(df, time_col, event_col, high_col):
    d = df[[time_col, event_col, high_col]].copy()
    d.columns = ["time", "event", "high"]
    d["time"] = pd.to_numeric(d["time"], errors="coerce")
    d["event"] = pd.to_numeric(d["event"], errors="coerce")
    d["high"] = pd.to_numeric(d["high"], errors="coerce")
    d = d.replace([np.inf, -np.inf], np.nan).dropna()
    d = d[(d["time"] > 0) & (d["event"].isin([0, 1]))]
    d["event"] = d["event"].astype(int)
    d["high"] = d["high"].astype(int)
    cph = CoxPHFitter()
    cph.fit(d, duration_col="time", event_col="event")
    hr = float(np.exp(cph.params_["high"]))
    l = float(np.exp(cph.confidence_intervals_.loc["high", "95% lower-bound"]))
    u = float(np.exp(cph.confidence_intervals_.loc["high", "95% upper-bound"]))
    p = float(cph.summary.loc["high", "p"])
    return f"{hr:.2f} ({l:.2f}–{u:.2f})", fmt_p(p)

fig7 = pd.read_csv(fig7_file)
cut_map = dict(zip(fig7["Cohort"].astype(str), fig7["Cutoff_Value"].astype(float)))
p_map = dict(zip(fig7["Cohort"].astype(str), fig7["LogRank_P"].astype(float)))
n_map = dict(zip(fig7["Cohort"].astype(str), fig7["N"].astype(int)))

tcga = pd.read_csv(risk_file)
tcga["high"] = (pd.to_numeric(tcga["Risk_Score"], errors="coerce") > float(cut_map["TCGA_train"])).astype(int)
hr_tcga, p_tcga = cox_hr_ci_p(tcga, "OS_Time", "OS_Event", "high")

rng = np.random.default_rng(2026)
idx = np.arange(len(tcga))
rng.shuffle(idx)
ntr = int(np.floor(0.7 * len(tcga)))
tcga_split_train = tcga.iloc[idx[:ntr]].copy()
tcga_split_val = tcga.iloc[idx[ntr:]].copy()
cut_split = float(np.median(pd.to_numeric(tcga_split_train["Risk_Score"], errors="coerce")))
tcga_split_val["high"] = (pd.to_numeric(tcga_split_val["Risk_Score"], errors="coerce") > cut_split).astype(int)
hr_split, p_split = cox_hr_ci_p(tcga_split_val, "OS_Time", "OS_Event", "high")

g39279 = pd.read_csv(g39279_file)
score39279 = "Risk_Score_Correct" if "Risk_Score_Correct" in g39279.columns else "risk_score"
g39279["high"] = (pd.to_numeric(g39279[score39279], errors="coerce") > float(cut_map["GSE39279_RFS"])).astype(int)
hr_39279, p_39279 = cox_hr_ci_p(g39279, "rfs_time", "rfs_status", "high")

g30219 = pd.read_csv(g30219_file)
g30219["high"] = (pd.to_numeric(g30219["MFAP2_Expr"], errors="coerce") > float(cut_map["GSE30219_OS"])).astype(int)
hr_30219, p_30219 = cox_hr_ci_p(g30219, "Time_Months", "Status", "high")

table5_rows = [
    ["TCGA (training)", "Overall survival (OS)", str(int(n_map["TCGA_train"])), hr_tcga, fmt_p(p_map["TCGA_train"]), "0.664", "0.685", "1.227", "0.843", "Median cutoff derived from the training cohort"],
    ["TCGA (internal validation)", "Overall survival (OS)", str(int(n_map["TCGA_internal_split"])), hr_split, fmt_p(p_map["TCGA_internal_split"]), "Not available", "Not available", "Not available", "Not available", "Median cutoff derived from the training cohort"],
    ["GSE39279 (external)", "Recurrence-free survival (RFS)", str(int(n_map["GSE39279_RFS"])), hr_39279, fmt_p(p_map["GSE39279_RFS"]), "Not available", "Not available", "Not available", "Not available", "Maxstat-derived optimal cutoff (external cohorts)"],
    ["GSE30219 (external)", "Overall survival (OS)", str(int(n_map["GSE30219_OS"])), hr_30219, fmt_p(p_map["GSE30219_OS"]), "Not available", "Not available", "Not available", "Not available", "Maxstat-derived optimal cutoff (external cohorts)"],
]

header = "| **Dataset** | **Endpoint (OS / RFS)** | **N** | **HR (high vs low risk, Cox model, 95% CI)** | **P-value** | **C-index** | **AUC (3-year)** | **Calibration slope (3Y)** | **Calibration R² (3Y)** | **Cutoff method** |"
align = "|:--|:--:|--:|:--|:--:|:--:|:--:|:--:|:--:|:--|"
body = "\n".join([f"| {r[0]} | {r[1]} | {r[2]} | {r[3]} | {r[4]} | {r[5]} | {r[6]} | {r[7]} | {r[8]} | {r[9]} |" for r in table5_rows])
new_block = "\n".join([header, align, body])

def replace_md_table(path):
    if not os.path.exists(path):
        return
    txt = open(path, "r", encoding="utf-8").read()
    note = "*Figure/Table consistency note: all metrics in this table are aligned with the final M12 model outputs used in Figure 5–7 and Supplementary Figure S6. HR in this table represents high versus low risk groups; HR per SD is reported in Supplementary Table S8.*"
    rep = new_block + "\n\n" + note
    pat = re.compile(
        r"\| \*\*Dataset.*?\|\n\|:--\|:--:.*?(?:\n\| .*?\|)+\n\n\*Figure/Table consistency note:.*?\*",
        re.S
    )
    txt2 = pat.sub(rep, txt, count=1)
    with open(path, "w", encoding="utf-8") as f:
        f.write(txt2)

replace_md_table(md_file)
replace_md_table(md_file2)

doc = Document(docx_file)
if len(doc.tables) >= 5:
    t5 = doc.tables[4]
    header_cells = t5.rows[0].cells
    header_vals = [
        "Dataset",
        "Endpoint (OS / RFS)",
        "N",
        "HR (high vs low risk, Cox model, 95% CI)",
        "P-value",
        "C-index",
        "AUC (3-year)",
        "Calibration slope (3Y)",
        "Calibration R² (3Y)",
        "Cutoff method",
    ]
    for j, v in enumerate(header_vals):
        if j < len(header_cells):
            header_cells[j].text = v
    while len(t5.rows) > 1:
        t5._tbl.remove(t5.rows[-1]._tr)
    for r in table5_rows:
        cells = t5.add_row().cells
        for j, v in enumerate(r):
            cells[j].text = str(v)
doc.save(docx_file)

print("Updated Main Table 5 with TCGA internal validation row")
