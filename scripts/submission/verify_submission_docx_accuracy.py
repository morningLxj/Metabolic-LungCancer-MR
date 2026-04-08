import re
import pandas as pd
from docx import Document

base = "E:/GWAS"
main_doc = f"{base}/Final_Tables_SCI_Format/Main_Tables_Final.docx"
supp_doc = f"{base}/Final_Tables_SCI_Format/Supplementary_Tables_Final.docx"

main = Document(main_doc)
supp = Document(supp_doc)

main_tables = main.tables
supp_tables = supp.tables

coef = pd.read_csv(f"{base}/PDC/analysis/model_selection_14cpg/BestModel_Coefficients.csv")
fig7 = pd.read_csv(f"{base}/Final_Submission_Figures/Main_Figures_14CpG/Figure7_M12_Summary.csv")
s1 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S1.csv")
s2 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S2.csv")
s3 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S3.csv")
s4 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S4.csv")
s5 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S5.csv")
s6 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S6.csv")
s7 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S7.csv")
s8 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S8.csv")
s9 = pd.read_csv(f"{base}/Final_Tables_SCI_Format/Supplementary_Table_S9.csv")

def table_to_df(tbl):
    rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
    if not rows:
        return pd.DataFrame()
    h = rows[0]
    d = rows[1:]
    return pd.DataFrame(d, columns=h)

main_df4 = table_to_df(main_tables[3]) if len(main_tables) >= 4 else pd.DataFrame()
main_df5 = table_to_df(main_tables[4]) if len(main_tables) >= 5 else pd.DataFrame()

checks = []

checks.append(("main_doc_exists", True))
checks.append(("supp_doc_exists", True))
checks.append(("main_table_count>=5", len(main_tables) >= 5))
checks.append(("supp_table_count>=9", len(supp_tables) >= 9))

checks.append(("main_table4_rows_12", len(main_df4) == 12))
if len(main_df4) == 12:
    doc_ids = set(main_df4.iloc[:, 0].astype(str))
    src_ids = set(coef["Feature"].astype(str))
    checks.append(("main_table4_cpg_set_match", doc_ids == src_ids))
else:
    checks.append(("main_table4_cpg_set_match", False))

checks.append(("main_table5_rows_4", len(main_df5) == 4))
if len(main_df5) == 4 and "Dataset" in main_df5.columns:
    sup_map = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹⁻", "0123456789-")
    def parse_p_text(s):
        t = str(s).strip().replace(",", "")
        t = t.replace("P =", "").replace("p =", "").replace("P<", "<").replace("p<", "<").strip()
        if "×" in t and "10" in t:
            parts = t.split("×")
            try:
                m = float(re.sub(r"[^0-9.\-]", "", parts[0]))
                e_part = parts[1].strip()
                if "10" in e_part:
                    e = e_part.split("10", 1)[1].translate(sup_map)
                    e = int(e)
                    return m * (10 ** e)
            except Exception:
                pass
        m = re.search(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", t)
        return float(m.group(0)) if m else None
    def norm_ds(x):
        s = str(x).lower()
        if "internal" in s and "tcga" in s:
            return "TCGA_internal_split"
        if "tcga" in s:
            return "TCGA_train"
        if "39279" in s:
            return "GSE39279_RFS"
        if "30219" in s:
            return "GSE30219_OS"
        return str(x)
    ds = set(main_df5["Dataset"].astype(str).map(norm_ds))
    checks.append(("main_table5_dataset_match", ds == set(fig7["Cohort"].astype(str))))
    pcol = "KM P-value" if "KM P-value" in main_df5.columns else ("P-value" if "P-value" in main_df5.columns else None)
    p_doc = dict(zip(main_df5["Dataset"].astype(str).map(norm_ds), main_df5[pcol].astype(str))) if pcol else {}
    ok = True
    for _, r in fig7.iterrows():
        k = str(r["Cohort"])
        p_src = float(r["LogRank_P"])
        if k not in p_doc:
            ok = False
            continue
        p_v = parse_p_text(p_doc[k])
        if p_v is None:
            ok = False
            continue
        if abs(p_v - p_src) > 5e-4:
            ok = False
    checks.append(("main_table5_km_p_consistent", ok))
else:
    checks.append(("main_table5_dataset_match", False))
    checks.append(("main_table5_km_p_consistent", False))

supp_expected = [s1, s2, s3, s4, s5, s6, s7, s8, s9]
for i, df in enumerate(supp_expected, start=1):
    ddoc = table_to_df(supp_tables[i - 1]) if len(supp_tables) >= i else pd.DataFrame()
    checks.append((f"supp_table_S{i}_rowcount_match", len(ddoc) == len(df)))

all_ok = all(v for _, v in checks)
for k, v in checks:
    print(k, "PASS" if v else "FAIL")
print("overall", "PASS" if all_ok else "FAIL")
